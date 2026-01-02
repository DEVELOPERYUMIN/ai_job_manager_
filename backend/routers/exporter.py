
import os
import io
import datetime
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from ..database import SessionLocal
from ..models import Resume, InterviewQuestion, InterviewAnswer, User
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from ..schemas import ExportResponse

router = APIRouter()

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def _get_user(db: Session, user_id: int):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return user

@router.get("/{user_id}/docx", response_model=ExportResponse)
def export_docx(user_id: int, db: Session = Depends(get_db)):
    user = _get_user(db, user_id)

    doc = Document()
    doc.add_heading(f"{user.name}님의 Job Prep Report", level=0)
    doc.add_paragraph(f"생성 일시: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(" ")

    # 1) 자기소개서
    doc.add_heading("1. 자기소개서", level=1)
    resumes = db.query(Resume).filter(Resume.user_id == user_id).all()
    if not resumes:
        doc.add_paragraph("등록된 자기소개서가 없습니다.")
    for r in resumes:
        doc.add_paragraph(f"• Original:\n{r.original_text}")
        if r.edited_text:
            doc.add_paragraph(f"• Edited:\n{r.edited_text}")
            doc.add_paragraph(f"• Feedback:\n{r.feedback}")
        doc.add_paragraph("---")

    # 2) 면접 질문
    doc.add_heading("2. 면접 질문", level=1)
    questions = db.query(InterviewQuestion).filter(InterviewQuestion.user_id == user_id).all()
    if not questions:
        doc.add_paragraph("생성된 면접 질문이 없습니다.")
    for idx, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{idx}. {q.question_text}")
    doc.add_paragraph("---")

    # 3) 답변 및 평가
    doc.add_heading("3. 면접 답변 및 평가", level=1)
    answers = db.query(InterviewAnswer).join(
        InterviewQuestion, InterviewQuestion.id == InterviewAnswer.question_id
    ).filter(InterviewQuestion.user_id == user_id).all()
    if not answers:
        doc.add_paragraph("등록된 답변이 없습니다.")
    for idx, a in enumerate(answers, start=1):
        question = db.query(InterviewQuestion).filter(InterviewQuestion.id == a.question_id).first()
        doc.add_paragraph(f"{idx}. Q: {question.question_text}")
        doc.add_paragraph(f"   A: {a.answer_text}")
        if a.score is not None:
            doc.add_paragraph(f"   - Score: {a.score}")
            doc.add_paragraph(f"   - Feedback: {a.feedback}")
        doc.add_paragraph("---")

    os.makedirs("exported_docs", exist_ok=True)
    filename = f"report_user_{user_id}.docx"
    filepath = os.path.join("exported_docs", filename)
    doc.save(filepath)

    return {"filename": filename, "download_url": f"/exporter/download/docx/{filename}"}

@router.get("/download/docx/{filename}")
def download_docx(filename: str):
    file_path = os.path.join("exported_docs", filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(path=file_path, filename=filename,
                        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@router.get("/{user_id}/pdf", response_model=ExportResponse)
def export_pdf(user_id: int, db: Session = Depends(get_db)):
    user = _get_user(db, user_id)

    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 40

    p.setFont("Helvetica-Bold", 14)
    p.drawString(40, y, f"{user.name}님의 Job Prep Report")
    y -= 30
    p.setFont("Helvetica", 10)
    p.drawString(40, y, f"생성 일시: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    y -= 30

    # 1) 자기소개서
    p.setFont("Helvetica-Bold", 12)
    p.drawString(40, y, "1. 자기소개서")
    y -= 20
    resumes = db.query(Resume).filter(Resume.user_id == user_id).all()
    if not resumes:
        p.setFont("Helvetica", 10)
        p.drawString(40, y, "등록된 자기소개서가 없습니다.")
        y -= 20
    for r in resumes:
        p.setFont("Helvetica-Bold", 10)
        p.drawString(40, y, "Original:")
        y -= 15
        p.setFont("Helvetica", 10)
        for line in r.original_text.split("\n"):
            p.drawString(60, y, line)
            y -= 12
            if y < 50:
                p.showPage()
                y = height - 40
        if r.edited_text:
            p.setFont("Helvetica-Bold", 10)
            p.drawString(40, y, "Edited:")
            y -= 15
            p.setFont("Helvetica", 10)
            for line in r.edited_text.split("\n"):
                p.drawString(60, y, line)
                y -= 12
                if y < 50:
                    p.showPage()
                    y = height - 40
            p.setFont("Helvetica-Bold", 10)
            p.drawString(40, y, "Feedback:")
            y -= 15
            p.setFont("Helvetica", 10)
            for line in r.feedback.split("\n"):
                p.drawString(60, y, line)
                y -= 12
                if y < 50:
                    p.showPage()
                    y = height - 40
        y -= 20

    # 2) 면접 질문
    p.setFont("Helvetica-Bold", 12)
    p.drawString(40, y, "2. 면접 질문")
    y -= 20
    questions = db.query(InterviewQuestion).filter(InterviewQuestion.user_id == user_id).all()
    if not questions:
        p.setFont("Helvetica", 10)
        p.drawString(40, y, "생성된 면접 질문이 없습니다.")
        y -= 20
    for idx, q in enumerate(questions, start=1):
        p.setFont("Helvetica", 10)
        p.drawString(60, y, f"{idx}. {q.question_text}")
        y -= 15
        if y < 50:
            p.showPage()
            y = height - 40
    y -= 20

    # 3) 답변 및 평가
    p.setFont("Helvetica-Bold", 12)
    p.drawString(40, y, "3. 면접 답변 및 평가")
    y -= 20
    answers = db.query(InterviewAnswer).join(
        InterviewQuestion, InterviewQuestion.id == InterviewAnswer.question_id
    ).filter(InterviewQuestion.user_id == user_id).all()
    if not answers:
        p.setFont("Helvetica", 10)
        p.drawString(40, y, "등록된 답변이 없습니다.")
        y -= 20
    for idx, a in enumerate(answers, start=1):
        question = db.query(InterviewQuestion).filter(InterviewQuestion.id == a.question_id).first()
        p.setFont("Helvetica-Bold", 10)
        p.drawString(40, y, f"{idx}. Q: {question.question_text}")
        y -= 15
        p.setFont("Helvetica", 10)
        for line in a.answer_text.split("\n"):
            p.drawString(60, y, line)
            y -= 12
            if y < 50:
                p.showPage()
                y = height - 40
        if a.score is not None:
            p.setFont("Helvetica-Bold", 10)
            p.drawString(40, y, f"   - Score: {a.score}")
            y -= 15
            p.setFont("Helvetica", 10)
            for line in a.feedback.split("\n"):
                p.drawString(60, y, line)
                y -= 12
                if y < 50:
                    p.showPage()
                    y = height - 40
        y -= 20

    p.showPage()
    p.save()

    buffer.seek(0)
    os.makedirs("exported_pdfs", exist_ok=True)
    filename = f"report_user_{user_id}.pdf"
    filepath = os.path.join("exported_pdfs", filename)
    with open(filepath, "wb") as f:
        f.write(buffer.getvalue())

    return {"filename": filename, "download_url": f"/exporter/download/pdf/{filename}"}

@router.get("/download/pdf/{filename}")
def download_pdf(filename: str):
    file_path = os.path.join("exported_pdfs", filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(path=file_path, filename=filename, media_type='application/pdf')
